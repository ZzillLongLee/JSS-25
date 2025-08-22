import olefile
from Crypto.Cipher import ARC4
from docx import Document


def unlock_and_save_hwp_to_word(filepath, password, output_word_file):
    try:
        ole = olefile.OleFileIO(filepath)

        # Check if the file has EncryptedSummaryInfo stream
        if ole.exists('EncryptedSummaryInfo'):
            stream = ole.openstream('EncryptedSummaryInfo')
            encrypted_data = stream.read()

            # Decrypt the data using ARC4
            cipher = ARC4.new(password.encode('utf-8'))
            decrypted_data = cipher.decrypt(encrypted_data)

            # Save decrypted data to a temporary file
            temp_hwp_file = 'decrypted_document.hwp'
            with open(temp_hwp_file, 'wb') as unlocked_file:
                unlocked_file.write(decrypted_data)

            print("HWP document decrypted successfully.")

            # Load the decrypted HWP file and copy its content to a Word document
            doc = Document()
            doc.add_heading('Decrypted HWP Content', level=1)

            # Note: This assumes the decrypted data is in plain text
            # If the decrypted data is in a specific format, additional parsing may be needed
            decrypted_text = decrypted_data.decode('utf-8', errors='ignore')
            doc.add_paragraph(decrypted_text)

            # Save to Word file
            doc.save(output_word_file)
            print(f"Decrypted content saved to {output_word_file}.")
        else:
            print("No EncryptedSummaryInfo stream found. This file may not be password protected.")

        ole.close()
    except Exception as e:
        print(f"An error occurred: {e}")


if __name__ == '__main__':
    hwp_path = r'C:\Users\DTaQ\Desktop\My task\2024 personal project\HRS\5.BS2_(HRS)_v0.1_230308'
    word_path = r'C:\Users\DTaQ\Desktop\output_document.docx'
    # Usage
    unlock_and_save_hwp_to_word(hwp_path, '000', word_path)
